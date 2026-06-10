import os
from datetime import datetime
from fastapi import UploadFile
from pathlib import Path
from sqlalchemy.orm import Session

from app.db.dependencies import get_db_session
from app.db.session import SessionLocal

from .schema import Processing_Type
from app.service.openai.service import OpenaiService
from app.models.schema import Document, File, Processing_Request, Extracted_Result, Processing_Job
from app.service.document.schema import Processing_status, Processing_Type, Processing_Job_Status
from app.tasks.tasks import app
from app.core.minio import minio_client
from app.service.file.file import post_file


def extract_text_from_pdf(file_stream):
    from pypdf import PdfReader
    reader = PdfReader(stream=file_stream)

    page = reader.pages[0]
    return page.extract_text()


def extract_text_from_doc(file_stream):
    from docx import Document

    # file_bytes = file_path.read()

    # file_stream = io.BytesIO(file_bytes)

    doc = Document(file_stream)

    text = [para.text for para in doc.paragraphs]
    full_text = "\n".join(text)

    return full_text


def get_file_extension(file: UploadFile) -> str:
    if file.filename is None:
        raise ValueError("Uploaded file has no filename")

    return Path(file.filename).suffix.lstrip(".")


async def process_document(file: UploadFile, processing_type: Processing_Type, instructions: str, db_session: Session):

    file_name = get_file_extension(file)

    file_object = File(name=file.filename,
                       content_type=file.content_type, extension=file_name)
    db_session.add(file_object)
    db_session.commit()
    db_session.refresh(file_object)

    result = await post_file(file)
    print("result", result)
    document_object = Document(name=file.filename, file=file_object)

    db_session.add(document_object)
    db_session.flush()

    processing_request_object = Processing_Request(
        document_id=document_object.id, processing_type=processing_type, instructions=instructions, status=Processing_status.QUEUED)

    db_session.add_all(
        [file_object, document_object, processing_request_object])

    db_session.commit()

    # start_processing.delay(processing_request_id=processing_request_object.id)

    return {"processing_request_id": processing_request_object.id, "status": processing_request_object.status}


@app.task(
    bind=True,
    autoretry_for=(Exception,),
    retry_backoff=True,
    retry_kwargs={"max_retries": 3}
)
def start_processing(processing_request_id: int):

    db_session = SessionLocal()
    try:
        processing_request_statement = db_session.query(Processing_Request).where(
            Processing_Request.id == processing_request_id
        )

        processing_request_result = processing_request_statement.scalar()

        document_statement = db_session.query(Document).where(
            Document.id == processing_request_result.document_id
        )
        document_statement_result = document_statement.scalar()

        file_statement = db_session.query(File).where(
            File.id == document_statement_result.file_id)
        file_statement_result = file_statement.scalar()

        file_object = minio_client.fget_object(bucket_name=os.getenv(
            "MINIO_BUCKET", ""), object_name=str(file_statement_result.id), file_path="/tmpt/files")

        processing_type = processing_request_result.processing_type
        instructions = processing_request_result.instructions

        extracted_content = ""
        if file_statement_result.content_type == "application/pdf":
            extracted_content = extract_text_from_pdf(file_stream=file_object)
        elif file_statement_result.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_content = extract_text_from_doc(file_stream=file_object)

        if not extracted_content:
            return

        if processing_type == Processing_Type.DOCUMENT_SUMMARY:
            result = OpenaiService().generate_summary(
                content=extracted_content, instructions=instructions)
        elif processing_type == Processing_Type.INVOICE_EXTRACTION:
            result = OpenaiService().get_invoice_metadata(
                content=extracted_content, instructions=instructions)
        elif processing_type == Processing_Type.CONTRACT_METADATA:
            result = OpenaiService().get_contract_metadata(
                content=extracted_content, instructions=instructions)
        else:
            return

        if result:
            processing_request_result.status = Processing_status.COMPLETED
            db_session.add(processing_request_result)
            db_session.commit()

            extracted_result_object = Extracted_Result(processing_request_id=processing_request_result.id,
                                                       result_type=processing_type.value, content_json={'result': result}, confidence_score=1.0)

            processing_job_object = Processing_Job(processing_request_id=processing_request_result.id,
                                                   attempt_number=1, started_at=datetime.now())

            db_session.add([extracted_result_object, processing_job_object
                            ])
            db_session.commit()

    except Exception as e:
        print("Error:", e)
    finally:
        db_session.close()


async def get_document(id: int):
    return


async def get_documents():
    return


async def delete_document(id: int):
    return
